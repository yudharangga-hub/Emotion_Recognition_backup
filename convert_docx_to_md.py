#!/usr/bin/env python3
"""
Convert DOCX file to Markdown format
Requires: python-docx library
Install with: pip install python-docx
"""

from docx import Document
import os
import re

def convert_docx_to_md(docx_file, output_file=None):
    """
    Convert a DOCX file to Markdown format
    
    Args:
        docx_file: Path to input DOCX file
        output_file: Path to output MD file (optional, defaults to same name with .md)
    """
    
    # Set output filename
    if output_file is None:
        output_file = docx_file.replace('.docx', '.md')
    
    # Load the DOCX file
    print(f"Loading DOCX file: {docx_file}")
    try:
        doc = Document(docx_file)
    except FileNotFoundError:
        print(f"Error: File not found - {docx_file}")
        return False
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return False
    
    # Extract content
    md_content = []
    
    print("Converting content...")
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            md_content.append("")
            continue
        
        # Detect heading levels based on style
        style_name = para.style.name if para.style else ""
        
        if 'Heading 1' in style_name or style_name == 'Heading 1':
            md_content.append(f"# {text}")
        elif 'Heading 2' in style_name or style_name == 'Heading 2':
            md_content.append(f"## {text}")
        elif 'Heading 3' in style_name or style_name == 'Heading 3':
            md_content.append(f"### {text}")
        elif 'Heading 4' in style_name or style_name == 'Heading 4':
            md_content.append(f"#### {text}")
        elif 'Heading 5' in style_name or style_name == 'Heading 5':
            md_content.append(f"##### {text}")
        elif 'Heading 6' in style_name or style_name == 'Heading 6':
            md_content.append(f"###### {text}")
        else:
            # Regular paragraph
            md_content.append(text)
    
    # Extract tables if any
    print(f"Found {len(doc.tables)} table(s)")
    
    for table in doc.tables:
        md_content.append("\n")
        
        # Extract table content
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            md_content.append("| " + " | ".join(cells) + " |")
        
        md_content.append("\n")
    
    # Join content
    final_content = "\n".join(md_content)
    
    # Clean up multiple blank lines
    final_content = re.sub(r'\n\n\n+', '\n\n', final_content)
    
    # Write to markdown file
    print(f"Writing to: {output_file}")
    try:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(final_content)
        print(f"✓ Successfully converted to Markdown!")
        print(f"Output file: {output_file}")
        return True
    except Exception as e:
        print(f"Error writing file: {e}")
        return False

def main():
    # Define file path
    docx_file = r"e:\Tesis\Emotion_Recognition_backup\DwiJokoNurdadi_TesisEmotion_rev.docx"
    md_file = r"e:\Tesis\Emotion_Recognition_backup\DwiJokoNurdadi_TesisEmotion_rev.md"
    
    # Check if file exists
    if not os.path.exists(docx_file):
        print(f"Error: File not found at {docx_file}")
        return
    
    print("=" * 60)
    print("DOCX to Markdown Converter")
    print("=" * 60)
    
    # Convert
    success = convert_docx_to_md(docx_file, md_file)
    
    if success:
        print("\n✓ Conversion completed successfully!")
        print(f"You can now read the file: {md_file}")
    else:
        print("\n✗ Conversion failed!")

if __name__ == "__main__":
    main()
